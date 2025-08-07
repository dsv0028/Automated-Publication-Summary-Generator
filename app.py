from flask import Flask, render_template, request, redirect, url_for, send_file, flash, jsonify, session
import pandas as pd
import json
from io import BytesIO
import os
import bibtexparser
from docx import Document
import google.generativeai as genai
from scholarly import scholarly
import concurrent.futures
import matplotlib.pyplot as plt
import seaborn as sns
from flask_cors import CORS
database = {}
app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})
DATABASE_FILE = 'database.json'
def load_database():
    if os.path.exists(DATABASE_FILE):
        with open(DATABASE_FILE, 'r') as f:
            return json.load(f)
    return {}

def save_database():
    with open(DATABASE_FILE, 'w') as f:
        json.dump(database, f)
app.secret_key = 'supersecretkey'  # Ensure you have a secret key for sessions and flash messages
database = load_database() 
print("Current database:",database)
# Initialize Google Generative AI
genai.configure(api_key="AIzaSyA_CXDGqn0nmoLeKd60yUg8l7F9uL3PSgk") # Replace with your actual API key
model = genai.GenerativeModel('gemini-1.5-flash-latest')

# Global variables
processed_df = None  # Placeholder for DataFrame
results_html = None  # Placeholder for HTML table
@app.route('/instructions')
def instructions():
    return render_template('instructions.html')


@app.route('/register')
def register():
    return render_template('registration.html')

@app.route('/author_prompt', methods=['POST'])
def author_prompt():
    if request.form['is_author'] == 'yes':
        return render_template('login.html')  # Show login if they are an author
    else:
        return render_template('index.html')  # Redirect to index if not an author



@app.route('/form_register', methods=['POST'])
def form_register():
    new_username = request.form['username']
    new_password = request.form['password']

    if new_username in database:
        # If username exists, inform the user
        return render_template('registration.html', info="Username already exists. Please login or choose a different username.")
    else:
        # If username is unique, add it to the database and redirect to login page
        database[new_username] = new_password
        save_database()
        print("Updated Database:",database)
        return render_template('login.html', info="Registration successful! Please log in.")
# Function to analyze author data
def d_analysis(author_name):
  search_query = scholarly.search_author(author_name)
  first_author_result = next(search_query)
  author = scholarly.fill(first_author_result)
  years = []
  for i in author['publications']:
    years.append(i['bib'].get('pub_year'))
  yy = [year for year in years if year is not None]

  hindex = author['hindex']
  iindex = author['i10index']
  year_count = {}


  for year in yy:
      if year in year_count:
          year_count[year] += 1
      else:
          year_count[year] = 1

  ordered_year_count = dict(sorted(year_count.items()))
  jif=[]
  yearlist=[]
  for key, val in author['cites_per_year'].items():
    sum = 0
    for kk, val in ordered_year_count.items():
      if(int(kk)==key-1 or int(kk)==key-2):
        sum = sum + val
      elif(int(kk)>key-1):
        break
    if(sum==0):
      continue
    jif.append(val/sum)
    yearlist.append(key)

    d_a = pd.DataFrame()
    d_a['Author'] = [author_name]*len(jif)
    d_a['Year'] = yearlist
    d_a['JIF'] = jif
    d_a['H-index'] = [hindex]*len(jif)
    d_a['I-index'] = [iindex]*len(jif)


  return d_a

# Function to retrieve scholarly information
def retrieve_stuffs(author_name, institution_name):
    try:
        search_query = f"{author_name} {institution_name}"
        search_result = scholarly.search_author(search_query)
        first_author_result = next(search_result)
        author = scholarly.fill(first_author_result)

        titles = [pub['bib']['title'] for pub in author['publications']]
        years = [pub['bib'].get('pub_year') for pub in author['publications']]
        citation = [pub['bib'].get('citation', '').lower() for pub in author['publications']]

        df = pd.DataFrame({
            'Author': [author_name] * len(author['publications']),
            'Title': titles,
            'Publication Year': years,
            'Citation': citation,
            'Institution_Name': institution_name
        })

        return df
    except StopIteration:
        flash(f"Author '{author_name}' not found. Please check the name and try again.")
        return pd.DataFrame()
    except Exception as e:
        flash(f"Error processing {author_name}: {str(e)}")
        return pd.DataFrame()


def generate_author_summary(df, author):
    filtered_df = df[df['Author'] == author]
    if filtered_df.empty:
        return f"No publications found for author '{author}'"

    titles = filtered_df['Title'].tolist()
    citations = filtered_df['Citation'].tolist()
    intro_phrases = [
        "One of the key works was",
        "Another notable publication was",
        "Among the significant contributions was",
        "An important study was",
        "A remarkable work was",
        "Additionally, there was",
        "Another critical work was"
    ]
    summary = f"{author} made significant contributions to their field with several impactful publications. "
    for i, (title, citation) in enumerate(zip(titles, citations)):
        phrase = intro_phrases[i % len(intro_phrases)]
        summary += f"{phrase} '{title}', which appeared in {citation}. "
    summary += f"These publications underscore {author}'s commitment to advancing research in their domain, particularly in areas such as {', '.join([title.split(':')[0].lower() for title in titles[:2]])}, and other related fields."

    SYSTEM_PROMPT = "Your name is Summarize AI. Your task is to Summarize the context. Aim for approximately 500 words."

    chat = model.start_chat(history=[{"role": "model", "parts": [SYSTEM_PROMPT]}])
    response = chat.send_message(f"""
User Prompt: Summarize the following text.
\n\n
Here is the Text: {summary}
\n\n
Instruction:
Generate a concise summary of the provided text.
""", stream=True)

    final_summary = ""
    for chunk in response:
        final_summary += chunk.text

    return final_summary

def generate_word_doc(summary_text):
    doc = Document()
    doc.add_heading('Customized Summary', level=1)
    doc.add_paragraph(summary_text)
    return doc

@app.route('/')
def hello_world():
    return render_template("login.html")
 
@app.route('/index')
def index_home():
    return render_template("index.html")



# Login logic
@app.route('/form_login', methods=['POST', 'GET'])
def login():
    name1 = request.form['username']
    pwd = request.form['password']

    if not database:
        # If the database is empty, show 'Not registered' message
        return render_template('login.html', info="No users are registered. Please register first.")
    
    if name1 not in database:
        return render_template('login.html', info='Invalid User - Not Registered')
    elif database[name1] != pwd:
        return render_template('login.html', info='Invalid Password')
    else:
        return render_template('index.html')

@app.route('/search', methods=['POST'])
def search_by_author():
    try:
        author_name = request.form.get('author_name')
        if not author_name:
            flash('Please enter an author name.')
            return redirect(request.url)

        # You may define institution_name or fetch it from a default source
        institution_name = ''  # Set to default or from a database

        # Retrieve and process data for the author
        result = retrieve_stuffs(author_name, institution_name)

        global processed_df, results_html
        processed_df = result
        results_html = result.to_html(classes='table dataTable', index=False)

        return redirect(url_for('results'))

    except Exception as e:
        flash(f'An error occurred during the search: {str(e)}')
        return redirect(url_for('index'))


@app.route('/upload', methods=['POST'])
def upload():
    if 'file' not in request.files:
        flash('No file part. Please select a file to upload.')
        return redirect(request.url)

    file = request.files['file']

    if file.filename == '':
        flash('No selected file. Please choose a valid file.')
        return redirect(request.url)

    if file:
        try:
            # Determine file type
            file_extension = os.path.splitext(file.filename)[1].lower()

            if file_extension == '.csv':
                df = pd.read_csv(file)
            elif file_extension == '.xlsx':
                df = pd.read_excel(file)
            elif file_extension == '.bib':
                # Read BibTeX file
                bib_data = bibtexparser.load(file)
                # Extract data into DataFrame (modify as needed)
                authors = [entry['author'] for entry in bib_data.entries]
                institutions = [entry.get('institution', '') for entry in bib_data.entries]
                df = pd.DataFrame({'Author': authors, 'Institution_Name': institutions})
            else:
                flash('Unsupported file format. Please upload a .csv, .xlsx, or .bib file.')
                return redirect(request.url)

            combined_df = pd.DataFrame()

            institution_name = df['Institution_Name'].iloc[0] if 'Institution_Name' in df.columns else ''
            author_list = df['Author'].tolist()

            with concurrent.futures.ThreadPoolExecutor() as executor:
                results = list(executor.map(lambda name: retrieve_stuffs(name, institution_name), author_list))

            for result in results:
                combined_df = pd.concat([combined_df, result], ignore_index=True)

            global processed_df, results_html
            processed_df = combined_df
            results_html = combined_df.to_html(classes='table dataTable', index=False)

            return redirect(url_for('results'))
                
        except Exception as e:
            flash(f'An error occurred while processing the file: {str(e)}')
            return redirect(url_for('index'))

@app.route('/results')
def results():
    if processed_df is None:
        flash("No results to display. Please upload and process a CSV first.")
        return redirect(url_for('index'))
    
    global results_html
    results_html = processed_df.to_html(classes='table dataTable', index=False)
    return render_template('results.html', results=results_html, authors=processed_df['Author'].unique().tolist())

@app.route('/summary', methods=['POST'])
def summary():
    author = request.form.get('author')
    if author:
        summary_text = generate_author_summary(processed_df, author)
        session['summary_text'] = summary_text
        return jsonify({'summary': summary_text})
    else:
        return jsonify({'summary': 'No author selected.'})

@app.route('/download_summary', methods=['GET'])
def download_summary():
    author = request.args.get('author')
    if 'summary_text' in session:
        summary_text = session['summary_text']  # Retrieve the summary from session
        summary_doc = generate_word_doc(summary_text)
        summary_buffer = BytesIO()
        summary_doc.save(summary_buffer)
        summary_buffer.seek(0)
        return send_file(summary_buffer, as_attachment=True, download_name=f"{author}_summary.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    else:
        flash("No summary available for download.")
        return redirect(url_for('results'))

@app.route('/download')
def download():
    if processed_df is None:
        flash("No data to download. Please upload and process a CSV first.")
        return redirect(url_for('index'))

    csv_buffer = BytesIO()
    processed_df.to_csv(csv_buffer, index=False)
    csv_buffer.seek(0)
    
    return send_file(csv_buffer, as_attachment=True, download_name='results.csv', mimetype='text/csv')

@app.route('/view_analysis', methods=['POST'])
def view_analysis():
    global processed_df
    
    if processed_df is None:
        return jsonify({'analysis': 'No data available for analysis.'})

    # Get unique authors from the processed DataFrame
    authors = processed_df['Author'].unique().tolist()

    # Analysis DataFrame
    analysis_dfs = [d_analysis(author) for author in authors]
    combined_analysis_df = pd.concat(analysis_dfs, ignore_index=True)

    # Generate publication trend visualization
    plt.figure(figsize=(10, 6))
    sns.countplot(data=combined_analysis_df, x='Year', hue='Author')
    plt.title('Publication Trends by Year')
    plt.xlabel('Year')
    plt.ylabel('Number of Publications')
    plt.xticks(rotation=45)
    
    # Save the figure to a BytesIO object
    image_buffer = BytesIO()
    plt.savefig(image_buffer, format='png')
    plt.close()
    image_buffer.seek(0)
    
    return jsonify({
        'analysis': combined_analysis_df.to_html(classes='table dataTable', index=False),
        'chart_data': combined_analysis_df['Author'].value_counts().to_dict(),
        'image': image_buffer.getvalue().decode('latin1')  # Convert to a string for JSON
    })


if __name__ == '__main__':
    app.run(debug=True) 